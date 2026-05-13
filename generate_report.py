"""
Generates the final project report in both .docx (editable) and .pdf
(Turnitin-ready) formats from a single source of truth.

Run from the project root:
    python3 generate_report.py

Outputs:
    Report_Final.docx
    Report_Final.pdf
"""
from __future__ import annotations

import csv
import os
from pathlib import Path

# ---------------- Pull the latest evaluation numbers from the CSV --------------
ROOT = Path(__file__).parent
SUMMARY_CSV = ROOT / "eval_model_summary.csv"


def load_model_summary():
    rows = []
    with SUMMARY_CSV.open() as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append({
                "model": r["model"],
                "queries_total": int(r["queries_total"]),
                "queries_with_gt": int(r["queries_with_gt"]),
                "precision": float(r["precision_at_10_mean"]),
                "recall": float(r["recall_at_10_mean"]),
                "f1": float(r["f1_at_10_mean"]),
                "latency_mean": float(r["latency_mean_sec"]),
                "latency_median": float(r["latency_median_sec"]),
            })
    return rows


def load_non_empty_counts():
    """Per-model count of non-empty answers from eval_detailed_results.csv (best-effort)."""
    detailed = ROOT / "eval_detailed_results.csv"
    counts = {}
    if not detailed.exists():
        return counts
    with detailed.open() as f:
        reader = csv.DictReader(f)
        for r in reader:
            m = r.get("model", "")
            if not m:
                continue
            counts.setdefault(m, {"total": 0, "non_empty": 0, "errors": 0})
            counts[m]["total"] += 1
            try:
                tokens = float(r.get("output_tokens_est") or 0)
                if tokens > 0:
                    counts[m]["non_empty"] += 1
            except ValueError:
                pass
            if r.get("error"):
                counts[m]["errors"] += 1
    return counts


MODEL_ROWS = load_model_summary()
NON_EMPTY = load_non_empty_counts()


def model_line(row):
    ne = NON_EMPTY.get(row["model"], {}).get("non_empty")
    suffix = f", {ne}/{row['queries_total']} non-empty answers" if ne is not None else ""
    return (
        f"{row['model']}: P@10={row['precision']:.3f}, "
        f"R@10={row['recall']:.3f}, F1@10={row['f1']:.3f}, "
        f"mean latency={row['latency_mean']:.2f}s "
        f"(median {row['latency_median']:.2f}s){suffix}"
    )


# ---------------- Static content (single source of truth) ---------------------
TITLE = "Graph-Based Intelligent Movie Recommendation System Using RAG"
SUBTITLE = "Project Report"
COURSE = "CMPE 258: Deep Learning - Spring 2026"
INSTRUCTOR = "Prof. Kaikai Liu, Department of Software Engineering, San Jose State University"

TEAM_ID = "Group 19"
TRACK = "System / Application"
FOCUS_AREAS = (
    "Knowledge Graphs, Retrieval-Augmented Generation (RAG), Recommender Systems, "
    "Semantic Search, FAISS, Large Language Models, Evaluation-First AI Systems"
)
TEAM_MEMBERS = [
    ("Bharath Kumar A",                              "018221268", "bharathkumar.a@sjsu.edu"),
    ("Saripella Sriyavarma",                          "019130553", "sriyavarma.saripella@sjsu.edu"),
    ("Venkata Siva Sai Krishna Prasad Yedupati",      "018320835", "venkatasivasaikrishnaprasad.yedupati@sjsu.edu"),
]

GITHUB_URL = "https://github.com/abharathkumarr/Graph-Based-Intelligent-Movie-Recommendation-System-Using-RAG"

ABSTRACT = (
    "Traditional movie recommendation systems based on collaborative or content-based "
    "filtering struggle with cold-start, popularity bias, and a lack of interpretability. "
    "We present a Graph-Based Intelligent Movie Recommendation System that combines "
    "a Neo4j knowledge graph with a Retrieval-Augmented Generation (RAG) pipeline. "
    "We extract 229,894 subject-predicate-object triplets from the Neo4j 'recommendations' "
    "dataset, render them as natural-language sentences, embed them with "
    "SentenceTransformers (all-MiniLM-L6-v2), and index them in FAISS for sub-second "
    "semantic retrieval. A LangChain RetrievalQA chain feeds the top-k retrieved triplets "
    "to a Groq-hosted LLM, which produces fact-grounded answers with citations. "
    "We deliver an evaluation-first study on 51 natural-language queries across three "
    "Groq LLMs, with 15 Cypher-derived ground-truth queries, and report Precision@10, "
    "Recall@10, F1@10, latency, and rate-limit / context-window failure modes. The "
    "evaluation drives an explicit production choice (llama-3.1-8b-instant: F1@10 = "
    f"{MODEL_ROWS[0]['f1']:.2f} at {MODEL_ROWS[0]['latency_mean']:.1f}s mean latency) and "
    "exposes real-world deployment trade-offs that pure offline accuracy comparisons hide."
)


# ---------------- DOCX builder ------------------------------------------------
def build_docx(path: Path):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Base style: Times New Roman 11pt
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    def H(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
        return h

    def P(text, bold=False, italic=False, align=None, size=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if align:
            p.alignment = align
        return p

    def Bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)

    def Code(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def Img(rel_path, width_inches=5.5, caption=None):
        full = ROOT / rel_path
        if not full.exists() or rel_path.endswith(".html"):
            return
        doc.add_picture(str(full), width=Inches(width_inches))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(caption)
            r.italic = True
            r.font.size = Pt(10)

    # -------- Cover --------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(TITLE)
    tr.bold = True
    tr.font.size = Pt(20)

    P(SUBTITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    P("")
    P(COURSE, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    P(INSTRUCTOR, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    P("")
    P(f"Team ID: {TEAM_ID}", bold=True)
    P(f"Project Track: {TRACK}", bold=True)
    P(f"Focused Areas: {FOCUS_AREAS}")
    P("")

    P("Team Members:", bold=True)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Name"
    hdr[1].text = "SJSU ID"
    hdr[2].text = "Email"
    for name, sid, email in TEAM_MEMBERS:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = sid
        row[2].text = email
    P("")
    P(f"Code repository: {GITHUB_URL}", italic=True)

    doc.add_page_break()

    # -------- Abstract --------
    H("Abstract", level=1)
    P(ABSTRACT)

    # -------- 1. Introduction --------
    H("1. Introduction and Problem Description", level=1)
    P(
        "Movie recommendation is a high-impact application of machine learning, but "
        "classical solutions are dominated by collaborative filtering and content-based "
        "approaches that have well-known weaknesses: popularity bias, cold-start "
        "for new users or items, and a lack of explainability. As streaming catalogs "
        "grow into the hundreds of thousands of titles, users increasingly issue "
        "natural-language queries (\"a recent adventure film with Leonardo DiCaprio that "
        "is rated above 7\") that simple matrix-factorization models cannot answer."
    )
    P(
        "Our project addresses this with a hybrid Knowledge-Graph + Retrieval-Augmented "
        "Generation (KG-RAG) architecture. We model the movie domain as a Neo4j graph, "
        "convert its facts into a corpus of natural-language triplets, index them with "
        "FAISS for sub-second semantic retrieval, and let a large language model "
        "produce grounded answers from the retrieved facts. The target users are "
        "viewers asking conversational queries and researchers/engineers who need an "
        "explainable, fact-anchored alternative to opaque embedding-only retrievers."
    )
    P("Concrete contributions of this work:", bold=True)
    Bullet("A reproducible Neo4j -> 229,894-triplet -> FAISS -> RAG pipeline that "
           "answers natural-language movie queries with fact-grounded citations.")
    Bullet("A Streamlit-based interactive UI exposing the retrieved triplets, "
           "answer latency, lexical grounding score, and a per-query pipeline trace - "
           "all of which support evaluation-first transparency.")
    Bullet("An extended evaluation study on 51 natural-language queries across 3 Groq "
           "LLMs, with 15 Cypher-derived ground-truth labels. The study turns API "
           "rate-limit, context-window, and payload constraints into first-class "
           "evaluation findings rather than incidental bugs.")
    Bullet("An evaluation-driven model-selection rationale (production pick = "
           "llama-3.1-8b-instant) rather than a defaulting-to-the-largest-model heuristic.")

    # -------- 2. Related Work --------
    H("2. Background and Related Work", level=1)
    P(
        "Lewis et al. [1] introduced Retrieval-Augmented Generation (RAG), which "
        "couples a dense passage retriever with a generative LLM so that responses "
        "are conditioned on retrieved evidence. Hogan et al. [2] surveyed knowledge "
        "graphs as a structured way to represent inter-entity relationships, motivating "
        "their use as a retrieval substrate. KG-Retriever [3] proposed efficient "
        "knowledge indexing for RAG over knowledge graphs, demonstrating gains from "
        "structured retrieval. Nair and Cheriyan [4] applied particle filtering over "
        "multi-featured movie knowledge graphs to refine recommendations, showing "
        "that graph traversal can capture user preferences that flat collaborative "
        "filtering misses. Xu et al. [5] showed that KG-grounded RAG meaningfully "
        "improves factuality and answer specificity in customer-service question "
        "answering."
    )
    P(
        "Our system synthesizes these threads: we follow [1] for the RAG pattern, [2] "
        "for graph-structured knowledge representation, and [3, 5] for the importance "
        "of indexing graph facts for retrieval. Our key engineering choice is to "
        "linearize graph edges into natural-language sentences before embedding, "
        "which lets us reuse off-the-shelf sentence encoders and a flat FAISS index "
        "while still keeping the structural facts intact for the LLM to consume."
    )

    # -------- 3. Dataset --------
    H("3. Dataset", level=1)
    P(
        "We use the public Neo4j 'recommendations' dataset hosted at "
        "bolt://demo.neo4jlabs.com:7687. It contains 28,865 nodes and 166,262 "
        "relationships spanning the entities and edges below."
    )
    P("Node types:", bold=True)
    Bullet("Movie - title, year, IMDb rating, budget, revenue, runtime, language")
    Bullet("Actor / Director - name, role, birth/death year")
    Bullet("User - userId, rating history")
    Bullet("Genre - genre label used to classify movies")
    P("Relationship types:", bold=True)
    Bullet("(:Movie)-[:IN_GENRE]->(:Genre)")
    Bullet("(:User)-[:RATED]->(:Movie)")
    Bullet("(:Actor)-[:ACTED_IN]->(:Movie)")
    Bullet("(:Director)-[:DIRECTED]->(:Movie)")
    Bullet("(:Movie)-[:RELEASED]->(:Year)")
    P(
        "Relevance: the Neo4j graph provides a multi-relational structure connecting "
        "movies to genres, actors, directors, and users, which is inherently suited "
        "to multi-hop reasoning. Rather than executing explicit multi-hop Cypher "
        "queries at inference time, we collect single-hop facts in advance, render "
        "them as natural-language triplets, and let the LLM perform implicit multi-hop "
        "reasoning over the retrieved set."
    )

    # -------- 4. System Design --------
    H("4. System / Model / Algorithm Design", level=1)
    P("The pipeline has six logical stages:", bold=True)
    Bullet("Knowledge Graph Construction: connect to Neo4j and execute Cypher queries "
           "to enumerate Movie/Actor/Director/Genre/User edges.")
    Bullet("Readable Triplet Generation: convert each edge into a (subject, predicate, "
           "object) tuple and then a sentence, e.g. \"Toy Story was released in year 1995\".")
    Bullet("Embedding Generation: encode each sentence with SentenceTransformers "
           "(all-MiniLM-L6-v2) into a 384-dimensional dense vector.")
    Bullet("FAISS Index Construction: store all 229,894 embeddings in a flat L2 index "
           "for sub-second nearest-neighbor search.")
    Bullet("Retrieval with Re-ranking: retrieve the top-k FAISS neighbors for the "
           "query embedding, then re-rank with cosine similarity to prioritize semantically "
           "closest facts.")
    Bullet("RAG Chain: feed the top-k re-ranked triplets and the user question into a "
           "LangChain RetrievalQA chain that calls a Groq-hosted LLM and returns a "
           "fact-grounded natural-language answer.")

    Img("pipeline.png", width_inches=6.0,
        caption="Figure 1. End-to-end KG-RAG pipeline. The user query is embedded, "
                "matched against the FAISS-indexed triplet corpus, re-ranked by "
                "cosine similarity, and consumed by a Groq-hosted LLM to produce "
                "a fact-grounded answer.")

    # -------- 5. Implementation --------
    H("5. Implementation Details", level=1)
    P("Languages and frameworks:", bold=True)
    Bullet("Python 3.11 (data pipeline, evaluation), Cypher (Neo4j queries)")
    Bullet("LangChain + langchain-groq (RetrievalQA orchestration)")
    Bullet("sentence-transformers (all-MiniLM-L6-v2 embeddings)")
    Bullet("FAISS-CPU (IndexFlatL2 vector store)")
    Bullet("Neo4j Python driver (Bolt 7687)")
    Bullet("Streamlit (interactive demo UI)")
    Bullet("matplotlib + seaborn (evaluation plots)")
    P("Compute environment:", bold=True)
    Bullet("Development: Google Colab CPU/T4 runtime")
    Bullet("Embedding generation: CPU (one-shot, ~10 min); embeddings cached to "
           "triplet_embeddings.pkl (337 MB) and triplet_sentences.pkl (9.8 MB).")
    Bullet("Inference: Groq cloud LLM endpoints (free tier); local FAISS query "
           "latency is < 50 ms.")
    P("Important implementation decisions:", bold=True)
    Bullet("Per-model retrieval k: gemma2-class models use k=20 (8k context); "
           "llama-3.1-8b uses k=30 (free-tier payload cap); llama-3.3-70b uses k=40. "
           "Equal k across models was infeasible due to differing context windows "
           "and payload limits.")
    Bullet("Robust answer parsing: a regex/whitespace-tolerant extractor pulls a clean "
           "movie-title list out of every LLM's free-text answer so all three models "
           "are scored consistently.")
    Bullet("Exponential back-off + wall-clock budget per model (70B = 360s, 8B = 600s, "
           "gemma2 = 300s) to handle Groq free-tier 429s deterministically.")
    Bullet("Cypher-derived ground truth for 15 evaluation queries gives objective "
           "Precision/Recall/F1 (the remaining 36 queries are scored for non-empty "
           "answer rate and latency only).")
    Bullet("No credentials in code: GROQ_API_KEY is read exclusively from the "
           "environment; .env.example documents the variable, and .gitignore covers "
           "logs and local backups.")
    P("Code link: " + GITHUB_URL, italic=True)
    P("Key entry points in the repository:", bold=True)
    Bullet("CMPE258_Project_Code.ipynb - end-to-end notebook including the Option 2 "
           "extended evaluation cells")
    Bullet("app_full.py - Streamlit demo (multi-model selector, grounding score, "
           "pipeline trace, JSONL log)")
    Bullet("data/eval_queries.json - 51-query evaluation set")
    Bullet("evaluation_visualizations.py - matplotlib helpers for per-query, "
           "confusion-style, and Precision/Recall vs k plots")
    Bullet("eval_detailed_results.csv, eval_model_summary.csv, eval_partial_<model>.csv "
           "- reproducible result artifacts")

    # -------- 6. Task distribution --------
    H("6. Task Distribution and Contributions", level=1)
    P("Bharath Kumar A:", bold=True)
    P(
        "RAG pipeline architecture (LangChain RetrievalQA + Groq); FAISS retrieval "
        "workflow and embedding generation; design and implementation of the Option 2 "
        "extended evaluation (multi-model loop, rate-limit handling, wall-clock budget, "
        "robust title parser); evaluation tables, charts, and findings paragraph; "
        "Streamlit UI integration; report writing and rubric alignment."
    )
    P("Venkata Siva Sai Krishna Prasad Yedupati:", bold=True)
    P(
        "Neo4j knowledge-graph connection and Cypher-based data extraction; subject-"
        "predicate-object triplet generation and natural-language sentence rendering; "
        "Cypher-derived ground-truth construction for 15 evaluation queries; pipeline "
        "correctness verification (graph facts vs RAG output); system testing."
    )
    P("Saripella Sriyavarma:", bold=True)
    P(
        "UI/UX design for the Streamlit demo (knowledge-graph visualization background, "
        "model selector, results layout); experiment validation and screenshot capture; "
        "presentation slide deck; final documentation review."
    )
    P(
        "All team members jointly contributed to system integration, end-to-end testing, "
        "the live demo, and the final report review."
    )

    # -------- 7. Evaluation --------
    H("7. Evaluation and Testing Results", level=1)
    P(
        "Per the assignment rubric, evaluation is the most important section of this "
        "report. We report two complementary studies: a focused 2-query qualitative "
        "study that validates correctness, and an extended 51-query, 3-model study "
        "that drives model selection."
    )

    H("7.1 Metrics", level=2)
    Bullet("Precision@10 - proportion of the top-10 returned movies that match the "
           "Cypher-derived ground truth.")
    Bullet("Recall@10 - proportion of all ground-truth movies that appear in the top-10.")
    Bullet("F1@10 - harmonic mean of Precision@10 and Recall@10.")
    Bullet("Latency - end-to-end wall-clock time per query, including retrieval and LLM "
           "generation.")
    Bullet("Non-empty answer rate - fraction of queries that returned any text (a "
           "robustness metric exposing rate-limit / context-window failures).")
    Bullet("Estimated cost per 1k output tokens - documented per model.")

    H("7.2 Focused Qualitative Validation (Original 2 Queries)", level=2)
    P("Query A: \"Find movies acted by Leonardo DiCaprio released after 2000.\"", italic=True)
    Bullet("Cypher ground truth (14 titles): Gangs of New York (2002), Catch Me If You "
           "Can (2002), The Aviator (2004), Blood Diamond (2006), The Departed (2006), "
           "Revolutionary Road (2008), Inception (2010), Django Unchained (2012), "
           "The Wolf of Wall Street (2013), The Revenant (2015), etc.")
    Bullet("RAG output (production model llama-3.3-70b-versatile, k=500): 10/14 correct.")
    Bullet("Precision@10 = 0.70, Recall@10 = 0.50, F1@10 = 0.58.")
    P("Query B: \"Recent Adventure movie rated above 3 with Leonardo DiCaprio.\"", italic=True)
    Bullet("Cypher ground truth: The Revenant, Blood Diamond, The Beach.")
    Bullet("RAG output: The Revenant (2015), correctly identified with explanation. "
           "Match: yes.")
    P("These two queries validate that the pipeline retrieves the right structural "
      "facts and that the LLM grounds its answer in those facts.")

    H("7.3 Extended Evaluation: 51 Queries x 3 Models", level=2)
    P(
        "We evaluated three Groq-hosted LLMs - llama-3.3-70b-versatile, "
        "llama-3.1-8b-instant, and a smaller OSS model (openai/gpt-oss-20b, used "
        "as a substitute when gemma2-9b-it was decommissioned mid-experiment) - "
        "on the same 51-query set with the same retriever and prompt. 15 of those "
        "queries have Cypher-derived ground truth used for Precision/Recall/F1; "
        "the remaining 36 contribute latency and non-empty-answer-rate evidence."
    )

    # Results table
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "P@10"
    hdr[2].text = "R@10"
    hdr[3].text = "F1@10"
    hdr[4].text = "Mean Latency (s)"
    hdr[5].text = "Non-empty answers"
    for row in MODEL_ROWS:
        ne = NON_EMPTY.get(row["model"], {}).get("non_empty", "-")
        total = row["queries_total"]
        cells = tbl.add_row().cells
        cells[0].text = row["model"]
        cells[1].text = f"{row['precision']:.3f}"
        cells[2].text = f"{row['recall']:.3f}"
        cells[3].text = f"{row['f1']:.3f}"
        cells[4].text = f"{row['latency_mean']:.2f}"
        cells[5].text = f"{ne}/{total}"
    P("Table 1. Per-model averages across 51 queries (15 with Cypher ground truth). "
      "Non-empty answer count is the strongest robustness signal: it directly "
      "measures whether the model produced any output under free-tier constraints.",
      italic=True, size=10)

    Img("eval_model_quality_latency.png", width_inches=5.8,
        caption="Figure 2. Side-by-side per-model F1@10 and mean latency. "
                "llama-3.1-8b-instant Pareto-dominates - higher F1 and lower latency "
                "than both alternatives within the same free-tier budget.")
    Img("eval_llama_3_1_8b_instant_precision_recall_vs_k.png", width_inches=5.8,
        caption="Figure 3. Average Precision and Recall as a function of retrieval k "
                "for the production model. Precision is stable while Recall climbs "
                "monotonically, supporting our choice of k = 40 in production.")
    Img("eval_llama_3_1_8b_instant_per_query_metrics.png", width_inches=6.2,
        caption="Figure 4. Per-query Precision/Recall/F1@10 for the production model "
                "across the 15 ground-truth queries.")
    Img("eval_llama_3_1_8b_instant_confusion_matrix.png", width_inches=5.0,
        caption="Figure 5. Retrieval-style confusion matrix (relevant vs retrieved) "
                "aggregated over the 15 ground-truth queries for the production model.")

    H("7.4 Findings", level=2)
    P("Production pick:", bold=True)
    P(
        f"{MODEL_ROWS[0]['model']} is the production-recommended model. It achieves "
        f"the highest F1@10 ({MODEL_ROWS[0]['f1']:.3f}) AND the lowest mean latency "
        f"({MODEL_ROWS[0]['latency_mean']:.2f}s) while returning a non-empty answer for "
        f"{NON_EMPTY.get(MODEL_ROWS[0]['model'], {}).get('non_empty', '-')}/"
        f"{MODEL_ROWS[0]['queries_total']} queries. This is an evaluation-driven "
        "decision; the larger 70B model was not picked because, on the Groq free tier, "
        "rate limits constrained it to a 5/51 success rate during our run - any "
        "marginal quality gains it could have offered were dominated by its inability "
        "to actually answer the user."
    )
    P("Real-world deployment finding (a result, not a bug):", bold=True)
    P(
        "Pure offline accuracy comparisons would have ranked the 70B model first, "
        "but provider-side rate limits, context-window caps, and per-payload size "
        "caps materially reshape the design space. Treating those constraints as "
        "first-class evaluation evidence - via the wall-clock budget bail-out and "
        "non-empty-answer rate - is what turned our \"slow model\" problem into a "
        "concrete model-selection conclusion."
    )

    H("7.5 Reproduction", level=2)
    P("Steps to reproduce the evaluation results from a clean checkout:", bold=True)
    Code("# 1. Clone the repo and install dependencies\n"
         f"git clone {GITHUB_URL}.git\n"
         "cd Graph-Based-Intelligent-Movie-Recommendation-System-Using-RAG\n"
         "pip install -r requirements.txt\n\n"
         "# 2. Provide your Groq API key (no key is stored in the repo)\n"
         "export GROQ_API_KEY=\"<your_key>\"\n\n"
         "# 3. Open the notebook and run Section 'Option 2' cells\n"
         "jupyter notebook CMPE258_Project_Code.ipynb\n\n"
         "# 4. Verify generated artifacts match those in the repo\n"
         "ls eval_detailed_results.csv eval_model_summary.csv \\\n"
         "   eval_model_quality_latency.png \\\n"
         "   eval_llama_3_1_8b_instant_per_query_metrics.png\n")
    P("Correctness checks:", bold=True)
    Bullet("Each evaluation query's ground truth is generated by a deterministic "
           "Cypher query against the same Neo4j graph the RAG retriever indexed - "
           "so any disagreement is unambiguously a retrieval/LLM error, not a labeling "
           "error.")
    Bullet("eval_partial_<model>.csv is written after every model finishes so a "
           "rate-limit-induced partial run can still be inspected and audited.")
    Bullet("The Streamlit UI exposes the full pipeline trace (top-k triplets, the "
           "prompt sent to the LLM, the raw LLM answer, and a lexical grounding "
           "score) for any single query - allowing live, qualitative verification "
           "during the demo.")

    # -------- 8. UI screenshots --------
    H("8. User Interface and Demo", level=2)
    P(
        "We also ship a Streamlit-based interactive demo (app_full.py) that accepts "
        "natural-language queries, displays the answer, the grounding score, and "
        "the underlying retrieved triplets, and animates a knowledge-graph "
        "visualization as the page background. Model switching, latency, and the "
        "pipeline trace are exposed in the sidebar to support live evaluation-first "
        "demonstrations."
    )
    Img("assets/knowledge_graph_bg.png", width_inches=5.5,
        caption="Figure 6. Knowledge-graph background used by the Streamlit demo, "
                "rendered from the (:Actor)-[:ACTED_IN]->(:Movie) subgraph.")

    # -------- 9. Conclusion --------
    H("9. Conclusion and Future Work", level=1)
    P(
        "We presented an evaluation-first Knowledge-Graph + RAG movie recommendation "
        "system that converts 229,894 Neo4j edges into a FAISS-indexed natural-"
        "language corpus, retrieves the most relevant facts for each query, and "
        "generates fact-grounded answers via a Groq-hosted LLM. The 51-query x "
        "3-model evaluation backs up an explicit production choice - "
        "llama-3.1-8b-instant - on both quality and latency, and turns Groq free-"
        "tier rate-limit, context-window, and payload caps into legitimate "
        "deployment evidence."
    )
    P("Future work:", bold=True)
    Bullet("Hybrid retrieval that fuses dense FAISS retrieval with on-demand Neo4j "
           "Cypher traversal so the system can answer truly multi-hop queries that "
           "depend on transitively related facts.")
    Bullet("Personalization via per-user RATED history embedded as additional "
           "triplets and conditioned on session state.")
    Bullet("A learned re-ranker (cross-encoder) replacing the cosine re-rank for "
           "better top-k ordering.")
    Bullet("Expansion to larger graphs (IMDb, TMDB) and a multi-modal upgrade that "
           "incorporates posters and trailers.")

    # -------- 10. References --------
    H("10. References", level=1)
    refs = [
        "[1] P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive "
        "NLP Tasks,\" arXiv:2005.11401, 2020. https://arxiv.org/abs/2005.11401",
        "[2] A. Hogan et al., \"Knowledge Graphs,\" ACM Computing Surveys, vol. 54, "
        "no. 4, pp. 1-37, 2022. https://doi.org/10.1145/3447772",
        "[3] \"KG-Retriever: Efficient Knowledge Indexing for Retrieval-Augmented Large "
        "Language Models,\" arXiv:2412.05547, 2023. https://arxiv.org/html/2412.05547v1",
        "[4] L. S. Nair and J. Cheriyan, \"Multi-Featured Movie Recommendation Using "
        "Knowledge Graph,\" IDCIoT 2023. https://doi.org/10.1109/idciot56793.2023.10053435",
        "[5] Z. Xu et al., \"Retrieval-Augmented Generation with Knowledge Graphs for "
        "Customer Service Question Answering,\" SIGIR 2024. "
        "https://doi.org/10.1145/3626772.3661370",
        "[6] Neo4j, \"Example datasets - Getting Started,\" "
        "https://neo4j.com/docs/getting-started/appendix/example-data/",
        "[7] N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence Embeddings using "
        "Siamese BERT-Networks,\" EMNLP 2019. https://arxiv.org/abs/1908.10084",
        "[8] J. Johnson, M. Douze, H. Jegou, \"Billion-scale similarity search with "
        "GPUs,\" IEEE Trans. Big Data, 2019 (FAISS). https://arxiv.org/abs/1702.08734",
        "[9] LangChain, \"RetrievalQA Documentation,\" "
        "https://python.langchain.com/docs/modules/chains/popular/vector_db_qa",
        "[10] Groq, \"LLM Inference API Documentation,\" https://console.groq.com/docs",
    ]
    for r in refs:
        P(r)

    doc.save(str(path))
    print(f"Wrote {path}")


# ---------------- PDF builder -------------------------------------------------
def build_pdf(path: Path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
        PageBreak, KeepTogether,
    )
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

    doc = SimpleDocTemplate(
        str(path), pagesize=letter,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        topMargin=0.9 * inch, bottomMargin=0.8 * inch,
        title=TITLE, author=", ".join(n for n, _, _ in TEAM_MEMBERS),
    )

    ss = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=ss["Title"], fontSize=20, leading=24,
                                  alignment=TA_CENTER, spaceAfter=10)
    sub_style = ParagraphStyle("sub", parent=ss["Heading2"], fontSize=14, leading=18,
                                alignment=TA_CENTER, spaceAfter=8)
    h1 = ParagraphStyle("h1", parent=ss["Heading1"], fontSize=14, leading=18,
                         spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#10367a"))
    h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontSize=12, leading=16,
                         spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#10367a"))
    body = ParagraphStyle("body", parent=ss["BodyText"], fontSize=10.5, leading=14,
                           alignment=TA_JUSTIFY, spaceAfter=6)
    bullet = ParagraphStyle("bul", parent=body, leftIndent=18, bulletIndent=6,
                             spaceAfter=2)
    caption = ParagraphStyle("cap", parent=body, fontSize=9, leading=11,
                              alignment=TA_CENTER, textColor=colors.HexColor("#555555"),
                              italic=1, spaceAfter=10)
    code_style = ParagraphStyle("code", parent=body, fontName="Courier", fontSize=8.5,
                                 leading=11, leftIndent=10, textColor=colors.HexColor("#222222"),
                                 spaceAfter=8)
    italic = ParagraphStyle("ital", parent=body, fontName="Times-Italic")

    story = []

    def P(text, style=body):
        story.append(Paragraph(text, style))

    def B(text):
        P("- " + text, bullet)

    def Img(rel_path, width_inches=5.5, cap=None):
        full = ROOT / rel_path
        if not full.exists() or rel_path.endswith(".html"):
            return  # html background isn't includable in PDF
        try:
            img = Image(str(full), width=width_inches * inch,
                         height=width_inches * inch * 0.65, kind="proportional")
        except Exception:
            return
        story.append(img)
        if cap:
            P(cap, caption)
        else:
            story.append(Spacer(1, 6))

    def Code(text):
        from html import escape
        for line in text.splitlines():
            P(escape(line).replace(" ", "&nbsp;"), code_style)

    # Cover
    story.append(Paragraph(TITLE, title_style))
    story.append(Paragraph(SUBTITLE, sub_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph(COURSE, body))
    story.append(Paragraph(INSTRUCTOR, body))
    story.append(Spacer(1, 14))
    P(f"<b>Team ID:</b> {TEAM_ID}")
    P(f"<b>Project Track:</b> {TRACK}")
    P(f"<b>Focused Areas:</b> {FOCUS_AREAS}")
    story.append(Spacer(1, 8))
    P("<b>Team Members:</b>")
    rows = [["Name", "SJSU ID", "Email"]] + [list(t) for t in TEAM_MEMBERS]
    t = Table(rows, hAlign="LEFT", colWidths=[2.4 * inch, 1.1 * inch, 3.0 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#10367a")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 5),
        ("TOPPADDING", (0, 0), (-1, 0), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))
    P(f"Code repository: <a href=\"{GITHUB_URL}\">{GITHUB_URL}</a>", italic)

    story.append(PageBreak())

    P("Abstract", h1)
    P(ABSTRACT)

    P("1. Introduction and Problem Description", h1)
    P("Movie recommendation is a high-impact application of machine learning, but "
      "classical solutions are dominated by collaborative filtering and content-based "
      "approaches that have well-known weaknesses: popularity bias, cold-start "
      "for new users or items, and a lack of explainability. As streaming catalogs "
      "grow into the hundreds of thousands of titles, users increasingly issue "
      "natural-language queries (\"a recent adventure film with Leonardo DiCaprio that "
      "is rated above 7\") that simple matrix-factorization models cannot answer.")
    P("Our project addresses this with a hybrid Knowledge-Graph + Retrieval-Augmented "
      "Generation (KG-RAG) architecture. We model the movie domain as a Neo4j graph, "
      "convert its facts into a corpus of natural-language triplets, index them with "
      "FAISS for sub-second semantic retrieval, and let a large language model "
      "produce grounded answers from the retrieved facts. The target users are "
      "viewers asking conversational queries and researchers/engineers who need an "
      "explainable, fact-anchored alternative to opaque embedding-only retrievers.")
    P("<b>Concrete contributions of this work:</b>")
    B("A reproducible Neo4j -> 229,894-triplet -> FAISS -> RAG pipeline that "
      "answers natural-language movie queries with fact-grounded citations.")
    B("A Streamlit-based interactive UI exposing the retrieved triplets, "
      "answer latency, lexical grounding score, and a per-query pipeline trace.")
    B("An extended evaluation study on 51 queries across 3 Groq LLMs, with 15 "
      "Cypher-derived ground-truth labels and explicit handling of free-tier "
      "rate-limit, context-window, and payload constraints as evaluation evidence.")
    B("An evaluation-driven model-selection rationale "
      "(production pick = llama-3.1-8b-instant).")

    P("2. Background and Related Work", h1)
    P("Lewis et al. [1] introduced Retrieval-Augmented Generation (RAG), which "
      "couples a dense passage retriever with a generative LLM so that responses "
      "are conditioned on retrieved evidence. Hogan et al. [2] surveyed knowledge "
      "graphs as a structured way to represent inter-entity relationships, motivating "
      "their use as a retrieval substrate. KG-Retriever [3] proposed efficient "
      "knowledge indexing for RAG over knowledge graphs, demonstrating gains from "
      "structured retrieval. Nair and Cheriyan [4] applied particle filtering over "
      "multi-featured movie knowledge graphs to refine recommendations. Xu et al. [5] "
      "showed that KG-grounded RAG meaningfully improves factuality and answer "
      "specificity in customer-service question answering.")
    P("Our system synthesizes these threads: we follow [1] for the RAG pattern, [2] "
      "for graph-structured knowledge representation, and [3, 5] for the importance "
      "of indexing graph facts for retrieval. Our key engineering choice is to "
      "linearize graph edges into natural-language sentences before embedding, "
      "which lets us reuse off-the-shelf sentence encoders and a flat FAISS index "
      "while still keeping the structural facts intact for the LLM to consume.")

    P("3. Dataset", h1)
    P("We use the public Neo4j 'recommendations' dataset hosted at "
      "bolt://demo.neo4jlabs.com:7687. It contains 28,865 nodes and 166,262 "
      "relationships.")
    P("<b>Node types:</b>")
    B("Movie - title, year, IMDb rating, budget, revenue, runtime, language")
    B("Actor / Director - name, role, birth/death year")
    B("User - userId, rating history")
    B("Genre - genre label used to classify movies")
    P("<b>Relationship types:</b>")
    B("(:Movie)-[:IN_GENRE]-&gt;(:Genre)")
    B("(:User)-[:RATED]-&gt;(:Movie)")
    B("(:Actor)-[:ACTED_IN]-&gt;(:Movie)")
    B("(:Director)-[:DIRECTED]-&gt;(:Movie)")
    B("(:Movie)-[:RELEASED]-&gt;(:Year)")
    P("Relevance: the graph provides a multi-relational structure connecting "
      "movies to genres, actors, directors, and users, which is inherently suited "
      "to multi-hop reasoning. We collect single-hop facts in advance, render "
      "them as natural-language triplets, and let the LLM perform implicit multi-hop "
      "reasoning over the retrieved set.")

    P("4. System / Model / Algorithm Design", h1)
    P("The pipeline has six logical stages:")
    B("<b>Knowledge Graph Construction:</b> connect to Neo4j and execute Cypher queries "
      "to enumerate Movie/Actor/Director/Genre/User edges.")
    B("<b>Readable Triplet Generation:</b> convert each edge into a (subject, predicate, "
      "object) tuple and then a sentence (e.g. \"Toy Story was released in year 1995\").")
    B("<b>Embedding Generation:</b> encode each sentence with SentenceTransformers "
      "(all-MiniLM-L6-v2) into a 384-dimensional dense vector.")
    B("<b>FAISS Index Construction:</b> store all 229,894 embeddings in a flat L2 index.")
    B("<b>Retrieval with Re-ranking:</b> retrieve top-k FAISS neighbors and re-rank "
      "by cosine similarity to prioritize semantically closest facts.")
    B("<b>RAG Chain:</b> feed top-k re-ranked triplets and the question into a LangChain "
      "RetrievalQA chain backed by a Groq-hosted LLM.")
    Img("pipeline.png", width_inches=5.8,
        cap="Figure 1. End-to-end KG-RAG pipeline.")

    P("5. Implementation Details", h1)
    P("<b>Languages and frameworks:</b>")
    B("Python 3.11 (pipeline + evaluation), Cypher (Neo4j queries)")
    B("LangChain + langchain-groq (RetrievalQA orchestration)")
    B("sentence-transformers (all-MiniLM-L6-v2 embeddings)")
    B("FAISS-CPU (IndexFlatL2 vector store)")
    B("Streamlit (interactive demo UI), matplotlib + seaborn (plots)")
    P("<b>Compute environment:</b>")
    B("Development: Google Colab CPU/T4 runtime")
    B("Embedding generation: CPU one-shot (~10 min); cached to "
      "triplet_embeddings.pkl (337 MB) and triplet_sentences.pkl (9.8 MB).")
    B("Inference: Groq cloud LLM endpoints (free tier); local FAISS query "
      "latency &lt; 50 ms.")
    P("<b>Important implementation decisions:</b>")
    B("Per-model retrieval k (gemma2: 20, 8B: 30, 70B: 40) to respect each "
      "model's context window and payload cap.")
    B("Robust answer parsing extracts a clean movie-title list across all "
      "models for consistent scoring.")
    B("Exponential back-off + wall-clock budget per model (70B = 360s, 8B = 600s, "
      "gemma2 = 300s) to handle free-tier 429s deterministically.")
    B("Cypher-derived ground truth for 15 evaluation queries gives objective "
      "Precision/Recall/F1.")
    B("No credentials in code: GROQ_API_KEY is read exclusively from the "
      "environment; .env.example documents the variable.")
    P(f"<b>Code link:</b> <a href=\"{GITHUB_URL}\">{GITHUB_URL}</a>")

    P("6. Task Distribution and Contributions", h1)
    P("<b>Bharath Kumar A:</b> RAG pipeline architecture (LangChain RetrievalQA + Groq); "
      "FAISS retrieval workflow and embedding generation; design and implementation "
      "of the Option 2 extended evaluation (multi-model loop, rate-limit handling, "
      "wall-clock budget, robust title parser); evaluation tables, charts, and "
      "findings paragraph; Streamlit UI integration; report writing and rubric "
      "alignment.")
    P("<b>Venkata Siva Sai Krishna Prasad Yedupati:</b> Neo4j knowledge-graph "
      "connection and Cypher-based data extraction; triplet generation and "
      "natural-language sentence rendering; Cypher-derived ground-truth construction "
      "for 15 evaluation queries; pipeline correctness verification; system testing.")
    P("<b>Saripella Sriyavarma:</b> UI/UX design for the Streamlit demo "
      "(knowledge-graph background, model selector, results layout); experiment "
      "validation and screenshot capture; presentation slide deck; final "
      "documentation review.")
    P("All team members jointly contributed to system integration, end-to-end "
      "testing, the live demo, and the final report review.")

    P("7. Evaluation and Testing Results", h1)
    P("Per the assignment rubric, evaluation is the most important section of "
      "this report. We report two complementary studies: a focused 2-query "
      "qualitative study that validates correctness, and an extended 51-query "
      "3-model study that drives model selection.")

    P("7.1 Metrics", h2)
    B("Precision@10 - proportion of top-10 returned movies matching the Cypher-"
      "derived ground truth.")
    B("Recall@10 - proportion of all ground-truth movies appearing in the top-10.")
    B("F1@10 - harmonic mean of Precision@10 and Recall@10.")
    B("Latency - end-to-end wall-clock time per query.")
    B("Non-empty answer rate - fraction of queries returning any text (a "
      "robustness metric that exposes rate-limit / context-window failures).")

    P("7.2 Focused Qualitative Validation (Original 2 Queries)", h2)
    P("<i>Query A: \"Find movies acted by Leonardo DiCaprio released after 2000.\"</i>")
    B("Cypher ground truth (14 titles): Gangs of New York (2002), Catch Me If You "
      "Can (2002), The Aviator (2004), Blood Diamond (2006), The Departed (2006), "
      "Revolutionary Road (2008), Inception (2010), Django Unchained (2012), "
      "The Wolf of Wall Street (2013), The Revenant (2015), etc.")
    B("RAG output (production model llama-3.3-70b-versatile, k=500): 10/14 correct.")
    B("Precision@10 = 0.70, Recall@10 = 0.50, F1@10 = 0.58.")
    P("<i>Query B: \"Recent Adventure movie rated above 3 with Leonardo DiCaprio.\"</i>")
    B("Cypher ground truth: The Revenant, Blood Diamond, The Beach.")
    B("RAG output: The Revenant (2015), correctly identified with explanation. Match: yes.")

    P("7.3 Extended Evaluation: 51 Queries x 3 Models", h2)
    P("We evaluated three Groq-hosted LLMs - llama-3.3-70b-versatile, "
      "llama-3.1-8b-instant, and a smaller OSS model (openai/gpt-oss-20b, used as "
      "a substitute when gemma2-9b-it was decommissioned mid-experiment) - on the "
      "same 51-query set with the same retriever and prompt. 15 queries have "
      "Cypher-derived ground truth used for Precision/Recall/F1; the remaining 36 "
      "contribute latency and non-empty-answer-rate evidence.")

    # results table
    data = [["Model", "P@10", "R@10", "F1@10", "Mean Latency (s)", "Non-empty"]]
    for row in MODEL_ROWS:
        ne = NON_EMPTY.get(row["model"], {}).get("non_empty", "-")
        data.append([row["model"], f"{row['precision']:.3f}",
                     f"{row['recall']:.3f}", f"{row['f1']:.3f}",
                     f"{row['latency_mean']:.2f}", f"{ne}/{row['queries_total']}"])
    tbl = Table(data, colWidths=[2.0 * inch, 0.6 * inch, 0.6 * inch, 0.7 * inch,
                                   1.2 * inch, 0.9 * inch])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#10367a")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("ALIGN", (1, 1), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#eef3fb")]),
    ]))
    story.append(tbl)
    P("Table 1. Per-model averages across 51 queries (15 with Cypher ground truth). "
      "Non-empty answer count directly measures robustness under free-tier "
      "constraints.", caption)
    Img("eval_model_quality_latency.png", width_inches=5.4,
        cap="Figure 2. Per-model F1@10 and mean latency. llama-3.1-8b-instant "
            "Pareto-dominates - higher quality and lower latency at the same cost.")
    Img("eval_llama_3_1_8b_instant_precision_recall_vs_k.png", width_inches=5.4,
        cap="Figure 3. Average Precision and Recall vs retrieval k for the "
            "production model. Precision is stable; Recall climbs monotonically.")
    Img("eval_llama_3_1_8b_instant_per_query_metrics.png", width_inches=5.8,
        cap="Figure 4. Per-query Precision/Recall/F1@10 for the production model "
            "across the 15 ground-truth queries.")
    Img("eval_llama_3_1_8b_instant_confusion_matrix.png", width_inches=4.6,
        cap="Figure 5. Retrieval-style confusion matrix (relevant vs retrieved) "
            "for the production model.")

    P("7.4 Findings", h2)
    P("<b>Production pick:</b> "
      f"{MODEL_ROWS[0]['model']} is the production-recommended model. It achieves "
      f"the highest F1@10 ({MODEL_ROWS[0]['f1']:.3f}) AND the lowest mean latency "
      f"({MODEL_ROWS[0]['latency_mean']:.2f}s) while returning a non-empty answer for "
      f"{NON_EMPTY.get(MODEL_ROWS[0]['model'], {}).get('non_empty', '-')}/"
      f"{MODEL_ROWS[0]['queries_total']} queries.")
    P("<b>Real-world deployment finding (a result, not a bug):</b> Pure offline "
      "accuracy comparisons would have ranked the 70B model first, but provider-"
      "side rate limits, context-window caps, and per-payload size caps materially "
      "reshape the design space. Treating those constraints as first-class "
      "evaluation evidence - via the wall-clock budget bail-out and non-empty-"
      "answer rate - turned our \"slow model\" problem into a concrete "
      "model-selection conclusion.")

    P("7.5 Reproduction", h2)
    P("Steps to reproduce the evaluation results from a clean checkout:")
    Code(f"git clone {GITHUB_URL}.git\n"
         "cd Graph-Based-Intelligent-Movie-Recommendation-System-Using-RAG\n"
         "pip install -r requirements.txt\n"
         "export GROQ_API_KEY=\"<your_key>\"\n"
         "jupyter notebook CMPE258_Project_Code.ipynb\n"
         "# Run all cells of Section 'Option 2' to regenerate eval_*.csv and eval_*.png\n"
         "ls eval_detailed_results.csv eval_model_summary.csv\n")
    P("<b>Correctness checks:</b>")
    B("Each evaluation query's ground truth is generated by a deterministic Cypher "
      "query against the same Neo4j graph the RAG retriever indexed - so any "
      "disagreement is unambiguously a retrieval/LLM error, not a labeling error.")
    B("eval_partial_&lt;model&gt;.csv is written after every model finishes so a "
      "rate-limit-induced partial run can still be inspected.")
    B("The Streamlit UI exposes the full pipeline trace (top-k triplets, the prompt "
      "sent to the LLM, the raw LLM answer, and a lexical grounding score) for "
      "live qualitative verification during the demo.")

    P("8. User Interface and Demo", h2)
    P("The Streamlit demo (app_full.py) accepts natural-language queries, displays "
      "the answer, the grounding score, and the retrieved triplets, and animates a "
      "knowledge-graph visualization as the page background. Model switching, "
      "latency, and the pipeline trace are exposed in the sidebar to support "
      "live evaluation-first demonstrations.")
    Img("assets/knowledge_graph_bg.png", width_inches=5.0,
        cap="Figure 6. Knowledge-graph background used by the Streamlit demo, "
            "rendered from the (:Actor)-[:ACTED_IN]->(:Movie) subgraph.")

    P("9. Conclusion and Future Work", h1)
    P("We presented an evaluation-first Knowledge-Graph + RAG movie recommendation "
      "system that converts 229,894 Neo4j edges into a FAISS-indexed natural-"
      "language corpus, retrieves the most relevant facts for each query, and "
      "generates fact-grounded answers via a Groq-hosted LLM. The 51-query x "
      "3-model evaluation backs up an explicit production choice "
      "(llama-3.1-8b-instant) on both quality and latency, and turns Groq free-"
      "tier rate-limit, context-window, and payload caps into legitimate "
      "deployment evidence.")
    P("<b>Future work:</b>")
    B("Hybrid retrieval that fuses dense FAISS retrieval with on-demand Neo4j "
      "Cypher traversal for truly multi-hop queries.")
    B("Personalization via per-user RATED history embedded as additional triplets "
      "and conditioned on session state.")
    B("A learned cross-encoder re-ranker replacing the cosine re-rank.")
    B("Expansion to larger graphs (IMDb, TMDB) and multi-modal upgrades "
      "incorporating posters and trailers.")

    P("10. References", h1)
    refs = [
        "[1] P. Lewis et al., \"Retrieval-Augmented Generation for Knowledge-Intensive "
        "NLP Tasks,\" arXiv:2005.11401, 2020. "
        "<a href=\"https://arxiv.org/abs/2005.11401\">https://arxiv.org/abs/2005.11401</a>",
        "[2] A. Hogan et al., \"Knowledge Graphs,\" ACM Computing Surveys, vol. 54, "
        "no. 4, pp. 1-37, 2022. "
        "<a href=\"https://doi.org/10.1145/3447772\">https://doi.org/10.1145/3447772</a>",
        "[3] \"KG-Retriever: Efficient Knowledge Indexing for Retrieval-Augmented Large "
        "Language Models,\" arXiv:2412.05547, 2023. "
        "<a href=\"https://arxiv.org/html/2412.05547v1\">https://arxiv.org/html/2412.05547v1</a>",
        "[4] L. S. Nair and J. Cheriyan, \"Multi-Featured Movie Recommendation Using "
        "Knowledge Graph,\" IDCIoT 2023. "
        "<a href=\"https://doi.org/10.1109/idciot56793.2023.10053435\">https://doi.org/10.1109/idciot56793.2023.10053435</a>",
        "[5] Z. Xu et al., \"Retrieval-Augmented Generation with Knowledge Graphs for "
        "Customer Service Question Answering,\" SIGIR 2024. "
        "<a href=\"https://doi.org/10.1145/3626772.3661370\">https://doi.org/10.1145/3626772.3661370</a>",
        "[6] Neo4j, \"Example datasets - Getting Started,\" "
        "<a href=\"https://neo4j.com/docs/getting-started/appendix/example-data/\">https://neo4j.com/docs/getting-started/appendix/example-data/</a>",
        "[7] N. Reimers and I. Gurevych, \"Sentence-BERT: Sentence Embeddings using "
        "Siamese BERT-Networks,\" EMNLP 2019. "
        "<a href=\"https://arxiv.org/abs/1908.10084\">https://arxiv.org/abs/1908.10084</a>",
        "[8] J. Johnson, M. Douze, H. Jegou, \"Billion-scale similarity search with "
        "GPUs,\" IEEE Trans. Big Data, 2019 (FAISS). "
        "<a href=\"https://arxiv.org/abs/1702.08734\">https://arxiv.org/abs/1702.08734</a>",
        "[9] LangChain, \"RetrievalQA Documentation,\" "
        "<a href=\"https://python.langchain.com/docs/modules/chains/popular/vector_db_qa\">"
        "https://python.langchain.com/docs/modules/chains/popular/vector_db_qa</a>",
        "[10] Groq, \"LLM Inference API Documentation,\" "
        "<a href=\"https://console.groq.com/docs\">https://console.groq.com/docs</a>",
    ]
    for r in refs:
        P(r)

    doc.build(story)
    print(f"Wrote {path}")


if __name__ == "__main__":
    build_docx(ROOT / "Report_Final.docx")
    build_pdf(ROOT / "Report_Final.pdf")
    print("\nDone. Outputs:")
    print(f"  {ROOT / 'Report_Final.docx'}")
    print(f"  {ROOT / 'Report_Final.pdf'}")
